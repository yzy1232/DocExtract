"""
Word 文档处理器 - 使用 python-docx 解析 DOCX
"""
from typing import List
from app.processors.base_processor import BaseDocumentProcessor, DocumentParseResult, PageContent


class DocxProcessor(BaseDocumentProcessor):
    """DOCX 文档解析器"""

    @property
    def supported_mime_types(self) -> List[str]:
        return [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ]

    async def parse(self, file_content: bytes, filename: str) -> DocumentParseResult:
        import io
        import zipfile
        import xml.etree.ElementTree as ET
        from docx import Document
        from docx.oxml.ns import qn

        errors: List[str] = []
        paragraphs_text: List[str] = []
        tables: List[dict] = []
        metadata: dict = {}

        try:
            doc = Document(io.BytesIO(file_content))

            # 提取核心元数据
            try:
                props = doc.core_properties
                metadata = {
                    "title": props.title or "",
                    "author": props.author or "",
                    "created": str(props.created) if props.created else "",
                    "modified": str(props.modified) if props.modified else "",
                }
            except Exception:
                pass

            # 提取正文段落
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs_text.append(text)

            # 提取表格
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cell_texts = [cell.text.strip() for cell in row.cells]
                    rows.append(cell_texts)
                if rows:
                    headers = rows[0]
                    data_rows = rows[1:]
                    tables.append({"headers": headers, "rows": data_rows})

        except Exception as e:
            # python-docx 对少量结构不规范的文档较敏感，降级到 XML 直读以提高容错。
            try:
                with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                    xml_bytes = zf.read("word/document.xml")
                root = ET.fromstring(xml_bytes)
                namespaces = {
                    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                }
                texts = []
                for node in root.findall(".//w:t", namespaces):
                    if node.text:
                        texts.append(node.text)
                if texts:
                    paragraphs_text = ["".join(texts)]
                else:
                    errors.append(f"DOCX解析失败: {str(e)}")
            except Exception:
                errors.append(f"DOCX解析失败: {str(e)}")

        full_text = "\n".join(paragraphs_text)
        full_text = self._clean_text(full_text)
        language = self._detect_language(full_text[:2000])

        # DOCX 通常是单"虚拟页"，按段落数估算分页
        page_count = max(1, len(paragraphs_text) // 40)
        page = PageContent(
            page_number=1,
            raw_text=full_text,
            tables=tables,
            has_table=len(tables) > 0,
        )

        return DocumentParseResult(
            page_count=page_count,
            pages=[page],
            full_text=full_text,
            language=language,
            metadata=metadata,
            errors=errors,
        )


class ExcelProcessor(BaseDocumentProcessor):
    """XLSX 文档解析器"""

    @property
    def supported_mime_types(self) -> List[str]:
        return [
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ]

    async def parse(self, file_content: bytes, filename: str) -> DocumentParseResult:
        import io
        import zipfile
        import xml.etree.ElementTree as ET
        import openpyxl

        errors: List[str] = []
        pages: List[PageContent] = []

        try:
            wb = openpyxl.load_workbook(io.BytesIO(file_content), data_only=False)
            for sheet_idx, sheet_name in enumerate(wb.sheetnames, 1):
                ws = wb[sheet_name]
                rows = []
                texts = []
                for row in ws.iter_rows(values_only=True):
                    row_data = [str(cell) if cell is not None else "" for cell in row]
                    if any(cell for cell in row_data):
                        rows.append(row_data)
                        texts.append("\t".join(row_data))

                headers = rows[0] if rows else []
                data_rows = rows[1:] if len(rows) > 1 else []
                text = "\n".join(texts)

                pages.append(PageContent(
                    page_number=sheet_idx,
                    raw_text=text,
                    tables=[{"sheet": sheet_name, "headers": headers, "rows": data_rows}],
                    has_table=True,
                ))

        except Exception as e:
            errors.append(f"XLSX解析失败: {str(e)}")

        # 部分文件 workbook 元信息损坏时，openpyxl 可能返回空 sheet 列表；降级直接读取 worksheet XML。
        if not pages:
            try:
                def _get_ns(tag: str) -> str:
                    if tag.startswith("{") and "}" in tag:
                        return tag[1: tag.find("}")]
                    return "http://schemas.openxmlformats.org/spreadsheetml/2006/main"

                with zipfile.ZipFile(io.BytesIO(file_content)) as zf:
                    shared_strings: List[str] = []
                    if "xl/sharedStrings.xml" in zf.namelist():
                        ss_root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
                        ns = {"x": _get_ns(ss_root.tag)}
                        for si in ss_root.findall("x:si", ns):
                            text_parts = [t.text or "" for t in si.findall(".//x:t", ns)]
                            shared_strings.append("".join(text_parts))

                    sheet_files = sorted(
                        name for name in zf.namelist()
                        if name.startswith("xl/worksheets/") and name.endswith(".xml")
                    )

                    for sheet_idx, sheet_file in enumerate(sheet_files, 1):
                        root = ET.fromstring(zf.read(sheet_file))
                        ns = {"x": _get_ns(root.tag)}
                        rows: List[List[str]] = []
                        texts: List[str] = []

                        for row in root.findall(".//x:sheetData/x:row", ns):
                            row_data: List[str] = []
                            for cell in row.findall("x:c", ns):
                                cell_type = cell.get("t")
                                v_node = cell.find("x:v", ns)
                                inline_node = cell.find("x:is/x:t", ns)
                                formula_node = cell.find("x:f", ns)

                                value = ""
                                if cell_type == "s" and v_node is not None and v_node.text:
                                    try:
                                        idx = int(v_node.text)
                                        value = shared_strings[idx] if 0 <= idx < len(shared_strings) else ""
                                    except Exception:
                                        value = v_node.text or ""
                                elif inline_node is not None and inline_node.text:
                                    value = inline_node.text
                                elif formula_node is not None and formula_node.text:
                                    value = f"={formula_node.text}"
                                elif v_node is not None and v_node.text:
                                    value = v_node.text

                                row_data.append(str(value).strip())

                            if any(cell for cell in row_data):
                                rows.append(row_data)
                                texts.append("\t".join(row_data))

                        if rows:
                            headers = rows[0]
                            data_rows = rows[1:] if len(rows) > 1 else []
                            text = "\n".join(texts)
                            pages.append(PageContent(
                                page_number=sheet_idx,
                                raw_text=text,
                                tables=[{"sheet": sheet_file.rsplit("/", 1)[-1], "headers": headers, "rows": data_rows}],
                                has_table=True,
                            ))
            except Exception as e:
                if not errors:
                    errors.append(f"XLSX兜底解析失败: {str(e)}")

        full_text = "\n\n".join(p.raw_text for p in pages)
        return DocumentParseResult(
            page_count=len(pages),
            pages=pages,
            full_text=self._clean_text(full_text),
            language=self._detect_language(full_text[:2000]),
            errors=errors,
        )


class TextProcessor(BaseDocumentProcessor):
    """纯文本文档处理器"""

    @property
    def supported_mime_types(self) -> List[str]:
        return ["text/plain", "text/markdown", "text/x-markdown"]

    async def parse(self, file_content: bytes, filename: str) -> DocumentParseResult:
        # 自动检测编码
        import chardet
        detected = chardet.detect(file_content)
        encoding = detected.get("encoding") or "utf-8"
        try:
            text = file_content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            text = file_content.decode("utf-8", errors="replace")

        text = self._clean_text(text)
        page = PageContent(page_number=1, raw_text=text)

        return DocumentParseResult(
            page_count=1,
            pages=[page],
            full_text=text,
            language=self._detect_language(text[:2000]),
            encoding=encoding,
        )
